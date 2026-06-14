"""
services/extraction_service.py
استخراج البيانات المالية من المستندات — النسخة المحسّنة
"""

import os
import uuid
import logging
import asyncio
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, Dict, List, Tuple

from sqlalchemy.orm import Session

from claude_service import claude_service
from infrastructure import settings

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# الثوابت
# ─────────────────────────────────────────────
UPLOAD_DIR   = Path(os.getenv("UPLOAD_DIR", "/var/uploads/valuation"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE      = 20 * 1024 * 1024    # 20 MB
MAX_PROMPT_CHARS   = 12_000              # حد الـ prompt لـ Claude
CONFIDENCE_REVIEW  = 0.85               # أقل من هذا → يحتاج مراجعة

ALLOWED_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/csv",
}

EXTRACTION_SYSTEM = """أنت محلل مالي متخصص في استخراج البيانات من القوائم المالية.
مهمتك: قراءة المستند وإعادة البيانات المالية بصيغة JSON دقيقة.
- استخرج الأرقام كما هي بدون تقريب
- العملة الافتراضية SAR ما لم يُذكر غير ذلك
- إذا لم تجد قيمة، استخدم null وليس 0
- درجة الثقة: 1.0 = متأكد تماماً، 0.5 = تقدير، 0.0 = غير موجود"""


# ─────────────────────────────────────────────
# الدالة الرئيسية
# ─────────────────────────────────────────────
async def process_document(
    file_path: str,
    file_type: str,
    project_id: str,
    statement_type: str = "auto",
    db: Optional[Session] = None,
) -> Dict:
    """
    معالجة مستند مالي:
    1. التحقق  2. حفظ دائم  3. استخراج محتوى
    4. Claude  5. التحقق من البيانات  6. حفظ DB
    """
    file_id  = str(uuid.uuid4())
    tmp_path = Path(file_path)

    logger.info("بدء معالجة: %s [%s]", tmp_path.name, file_id)

    # ── 1. تحقق ──────────────────────────────
    _validate_file(tmp_path, file_type)

    # ── 2. نقل للتخزين الدائم ────────────────
    suffix      = tmp_path.suffix.lower()
    stored_path = UPLOAD_DIR / f"{file_id}{suffix}"
    shutil.move(str(tmp_path), str(stored_path))
    logger.info("📁 محفوظ: %s", stored_path)

    # ── 3. استخراج المحتوى ───────────────────
    try:
        # FIX 1: asyncio.to_thread — لا يعطّل event loop أثناء I/O
        document_content = await _extract_content(stored_path, file_type)
    except Exception as e:
        logger.error("فشل الاستخراج: %s", e)
        return _error_result(file_id, stored_path, f"فشل استخراج المحتوى: {e}")

    # ── 4. اكتشاف نوع القائمة ────────────────
    detected_type = (
        statement_type if statement_type != "auto"
        else _detect_statement_type(document_content["text"])
    )
    logger.info("نوع القائمة: %s", detected_type)

    # ── 5. استخراج بـ Claude ─────────────────
    extraction_result = await _extract_with_claude(
        document_content = document_content,
        statement_type   = detected_type,
        project_id       = project_id,
    )

    if not extraction_result["success"]:
        return _error_result(file_id, stored_path, extraction_result.get("error", "Claude error"))

    extracted_data = extraction_result["data"]

    # ── 6. التحقق من صحة البيانات ────────────
    validation = validate_extracted_data(extracted_data)
    confidence = _calculate_average_confidence(extracted_data)

    logger.info("ثقة الاستخراج: %.2f | أخطاء: %d", confidence, len(validation["errors"]))

    # ── 7. حفظ في DB ─────────────────────────
    if db:
        _save_to_db(
            db, file_id, project_id, stored_path,
            tmp_path.name, file_type, detected_type,
            extracted_data, confidence,
            document_content.get("page_count", 1),
            validation, extraction_result.get("message_id"),
        )

    return {
        "file_id":        file_id,
        "status":         "completed",
        "statement_type": detected_type,
        "storage_path":   str(stored_path),
        "page_count":     document_content.get("page_count", 1),
        "confidence":     confidence,
        "data":           extracted_data,
        "requires_review": (
            len(validation["errors"]) > 0 or confidence < CONFIDENCE_REVIEW
        ),
        "validation":     validation,
        "usage":          extraction_result.get("usage", {}),
    }


# ─────────────────────────────────────────────
# التحقق من الملف
# ─────────────────────────────────────────────
def _validate_file(path: Path, mime_type: str):
    if not path.exists():
        raise ValueError(f"الملف غير موجود: {path}")
    size = path.stat().st_size
    if size == 0:
        raise ValueError("الملف فارغ")
    if size > MAX_FILE_SIZE:
        raise ValueError(f"الحجم {size/1024/1024:.1f}MB يتجاوز 20MB")
    clean = mime_type.split(";")[0].strip().lower()
    if clean not in ALLOWED_MIME:
        raise ValueError(f"نوع غير مدعوم: {mime_type}")


# ─────────────────────────────────────────────
# استخراج المحتوى
# ─────────────────────────────────────────────
async def _extract_content(path: Path, mime_type: str) -> Dict:
    """يعيد dict يحتوي text + tables + page_count"""
    clean = mime_type.split(";")[0].strip().lower()

    # FIX 1: asyncio.to_thread — العمليات الثقيلة خارج event loop
    if clean == "application/pdf":
        return await asyncio.to_thread(extract_from_pdf, str(path))

    if clean in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        return await asyncio.to_thread(extract_from_excel, str(path))

    if clean == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return await asyncio.to_thread(_extract_word, str(path))

    if clean in ("text/plain", "text/csv"):
        text = path.read_text(encoding="utf-8", errors="replace")
        return {"type": "text", "text": text, "tables": [], "page_count": text.count("\n") + 1}

    raise ValueError(f"لا مستخرج لـ {mime_type}")


# ─────────────────────────────────────────────
# FIX 2: PDF — إغلاق الملف دائماً + استخراج الجداول
# ─────────────────────────────────────────────
def extract_from_pdf(file_path: str) -> Dict:
    """استخراج النص والجداول من PDF"""
    import fitz  # PyMuPDF — اسمه الصحيح fitz وليس pymupdf مباشرة

    # FIX 2a: context manager يضمن إغلاق الملف حتى عند الاستثناء
    with fitz.open(file_path) as doc:
        text_parts: List[str] = []
        tables: List[Dict]    = []

        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")

            # FIX 2b: تضمين النص الفارغ بعلامة واضحة بدلاً من تجاهله
            text_parts.append(
                f"\n--- صفحة {page_num + 1} ---\n"
                + (page_text if page_text.strip() else "[صفحة بدون نص — قد تكون صورة]")
            )

            # FIX 2c: استخراج الجداول مع معالجة استثناء per-page
            try:
                for tbl in page.find_tables():
                    raw = tbl.extract()
                    if raw:
                        tables.append({"page": page_num + 1, "data": raw})
            except Exception as e:
                logger.warning("تجاوز جدول في صفحة %d: %s", page_num + 1, e)

        page_count = len(doc)

    # FIX 2d: دمج نص الجداول مع النص الرئيسي ليصل إلى Claude
    table_text = _tables_to_text(tables)
    full_text  = "\n\n".join(text_parts)
    if table_text:
        full_text += f"\n\n=== جداول مستخرجة ===\n{table_text}"

    return {
        "type":       "pdf",
        "text":       full_text,
        "tables":     tables,
        "page_count": page_count,
    }


# ─────────────────────────────────────────────
# FIX 3: Excel — هيكل موحد + دعم header متعدد
# ─────────────────────────────────────────────
def extract_from_excel(file_path: str) -> Dict:
    """استخراج البيانات من Excel"""
    import pandas as pd

    # FIX 3a: engine صريح يتجنب تحذيرات openpyxl/xlrd
    engine = "openpyxl" if file_path.endswith(".xlsx") else "xlrd"
    xl = pd.ExcelFile(file_path, engine=engine)

    sheets_data: Dict = {}
    text_parts:  List[str] = []

    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name, header=None)

        # FIX 3b: تنظيف الصفوف الفارغة كلياً
        df.dropna(how="all", inplace=True)
        df.fillna("", inplace=True)

        sheets_data[sheet_name] = {
            "data":    df.to_dict(orient="records"),
            "columns": df.columns.tolist(),
            "shape":   list(df.shape),
        }

        # FIX 3c: تمثيل نصي للـ prompt
        text_parts.append(
            f"[ورقة: {sheet_name}]\n"
            + df.to_string(index=False, header=False, na_rep="—")
        )

    return {
        "type":       "excel",
        "text":       "\n\n".join(text_parts),
        "sheets":     sheets_data,          # للاستخدام البرمجي
        "tables":     [],
        "page_count": len(xl.sheet_names),
    }


def _extract_word(file_path: str) -> Dict:
    from docx import Document
    doc   = Document(file_path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
    text = "\n".join(lines)
    return {"type": "word", "text": text, "tables": [], "page_count": len(lines)}


def _tables_to_text(tables: List[Dict]) -> str:
    parts = []
    for t in tables:
        rows = [" | ".join(str(c or "") for c in row) for row in t["data"] if any(row)]
        if rows:
            parts.append(f"[جدول صفحة {t['page']}]\n" + "\n".join(rows))
    return "\n\n".join(parts)


# ─────────────────────────────────────────────
# اكتشاف نوع القائمة
# ─────────────────────────────────────────────
def _detect_statement_type(text: str) -> str:
    snippet = text[:800].lower()
    scores = {
        "income_statement": sum(1 for k in
            ["إيرادات","مبيعات","أرباح","خسائر","مصروفات","revenue","income","profit"]
            if k in snippet),
        "balance_sheet": sum(1 for k in
            ["أصول","خصوم","حقوق","assets","liabilities","equity","ميزانية"]
            if k in snippet),
        "cash_flow": sum(1 for k in
            ["تدفقات","نقدية","cash flow","تشغيلية","استثمارية","تمويلية"]
            if k in snippet),
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "unknown"


# ─────────────────────────────────────────────
# استخراج بـ Claude
# ─────────────────────────────────────────────
async def _extract_with_claude(
    document_content: Dict,
    statement_type: str,
    project_id: str,
) -> Dict:
    """يبني الـ prompt ويستدعي ClaudeService"""

    text   = document_content.get("text", "")
    tables = document_content.get("tables", [])

    # FIX 4: حد الـ prompt مع إشارة واضحة للاقتطاع
    snippet = text[:MAX_PROMPT_CHARS]
    if len(text) > MAX_PROMPT_CHARS:
        snippet += f"\n\n[... تم اقتصار النص — إجمالي: {len(text):,} حرف ...]"

    prompt = _build_extraction_prompt(snippet, statement_type, len(tables))

    return await claude_service.call(
        prompt         = prompt,
        system_prompt  = EXTRACTION_SYSTEM,
        response_format= "json",
        max_tokens     = 4096,
    )


def _build_extraction_prompt(text: str, statement_type: str, table_count: int) -> str:
    schemas = {
        "income_statement": '{"statement_type":"income_statement","currency":"SAR","periods":[{"year":2024,"revenue":null,"cost_of_revenue":null,"gross_profit":null,"operating_expenses":null,"ebitda":null,"depreciation_amortization":null,"ebit":null,"interest_expense":null,"net_profit":null,"earnings_per_share":null}],"line_items":{},"metadata":{},"confidence_score":0.9}',
        "balance_sheet":    '{"statement_type":"balance_sheet","currency":"SAR","periods":[{"year":2024,"total_assets":null,"current_assets":null,"cash_and_equivalents":null,"accounts_receivable":null,"inventory":null,"non_current_assets":null,"total_liabilities":null,"current_liabilities":null,"long_term_debt":null,"total_equity":null,"retained_earnings":null}],"line_items":{},"metadata":{},"confidence_score":0.9}',
        "cash_flow":        '{"statement_type":"cash_flow","currency":"SAR","periods":[{"year":2024,"operating_cash_flow":null,"investing_cash_flow":null,"financing_cash_flow":null,"capital_expenditure":null,"free_cash_flow":null,"net_change_in_cash":null,"beginning_cash":null,"ending_cash":null}],"line_items":{},"metadata":{},"confidence_score":0.9}',
        "unknown":          '{"statement_type":"unknown","currency":"SAR","periods":[],"line_items":{},"metadata":{},"confidence_score":0.5}',
    }
    schema = schemas.get(statement_type, schemas["unknown"])

    return f"""استخرج البيانات المالية وأعد JSON فقط (بدون أي نص خارجه):

{schema}

تعليمات:
- أرقام صحيحة بدون فواصل: 22000000
- استخرج كل السنوات المتاحة في periods[]
- line_items: بنود إضافية غير موجودة في الهيكل
- metadata: اسم الشركة، الفترة، المراجع إن وُجدت
- الجداول في المستند: {table_count}

المستند:
{text}"""


# ─────────────────────────────────────────────
# FIX 5: التحقق من صحة البيانات
# ─────────────────────────────────────────────
def validate_extracted_data(data: Dict) -> Dict:
    """
    تحقق منطقي من البيانات المالية المستخرجة.
    يعيد {"errors": [...], "warnings": [...]}
    """
    errors:   List[str] = []
    warnings: List[str] = []

    if not isinstance(data, dict):
        return {"errors": ["البيانات ليست كائن JSON صالح"], "warnings": []}

    periods = data.get("periods", [])
    if not periods:
        warnings.append("لم يُستخرج أي فترة مالية")

    for period in periods:
        year = period.get("year")

        # تحقق من قائمة الدخل
        if data.get("statement_type") == "income_statement":
            rev = period.get("revenue")
            gp  = period.get("gross_profit")
            np_ = period.get("net_profit")

            if rev is not None and gp is not None and gp > rev:
                errors.append(f"{year}: إجمالي الربح ({gp:,}) أكبر من الإيرادات ({rev:,})")

            if rev is not None and np_ is not None and abs(np_) > rev:
                warnings.append(f"{year}: صافي الربح/الخسارة يتجاوز الإيرادات — تحقق")

        # تحقق من الميزانية
        elif data.get("statement_type") == "balance_sheet":
            assets = period.get("total_assets")
            liab   = period.get("total_liabilities")
            equity = period.get("total_equity")

            if all(v is not None for v in [assets, liab, equity]):
                diff = abs(assets - (liab + equity))
                # FIX 5a: هامش 1% للفروق التقريبية
                if assets > 0 and diff / assets > 0.01:
                    errors.append(
                        f"{year}: الميزانية غير متوازنة — "
                        f"أصول {assets:,} ≠ خصوم {liab:,} + حقوق {equity:,}"
                    )

        # تحقق من التدفقات
        elif data.get("statement_type") == "cash_flow":
            op  = period.get("operating_cash_flow")
            inv = period.get("investing_cash_flow")
            fin = period.get("financing_cash_flow")
            net = period.get("net_change_in_cash")

            if all(v is not None for v in [op, inv, fin, net]):
                calculated = op + inv + fin
                if abs(calculated - net) > 1000:   # فارق أكثر من 1000
                    warnings.append(
                        f"{year}: صافي التدفق المحسوب {calculated:,} ≠ المُعلَن {net:,}"
                    )

    return {"errors": errors, "warnings": warnings}


# ─────────────────────────────────────────────
# حساب متوسط الثقة
# ─────────────────────────────────────────────
def _calculate_average_confidence(data: Dict) -> float:
    """
    FIX 6: متوسط مرجّح — يعاقب على الحقول null الكثيرة
    بدلاً من أخذ confidence_score من Claude فقط
    """
    base_confidence = float(data.get("confidence_score", 0.8))

    periods = data.get("periods", [])
    if not periods:
        return base_confidence * 0.5   # عقوبة: لا فترات

    null_ratios = []
    for period in periods:
        values      = [v for k, v in period.items() if k != "year"]
        null_count  = sum(1 for v in values if v is None)
        null_ratios.append(null_count / max(len(values), 1))

    avg_null   = sum(null_ratios) / len(null_ratios)
    completeness = 1.0 - avg_null

    # متوسط مرجّح: 70% ثقة Claude + 30% اكتمال البيانات
    return round(base_confidence * 0.7 + completeness * 0.3, 3)


# ─────────────────────────────────────────────
# حفظ في DB
# ─────────────────────────────────────────────
def _save_to_db(
    db, file_id, project_id, stored_path,
    original_name, mime_type, statement_type,
    extracted_data, confidence, page_count,
    validation, message_id,
):
    try:
        from models import FinancialDocument, DocumentStatus, StatementType
        doc = FinancialDocument(
            id                = file_id,
            project_id        = project_id,
            original_name     = original_name,
            storage_path      = str(stored_path),
            file_size         = stored_path.stat().st_size,
            mime_type         = mime_type,
            status            = DocumentStatus.completed,
            statement_type    = StatementType(statement_type),
            extracted_data    = extracted_data,
            confidence        = confidence,
            page_count        = page_count,
            validated         = len(validation["errors"]) == 0,
            validation_errors = validation["errors"],
            claude_message_id = message_id,
            processed_at      = datetime.now(timezone.utc),
        )
        db.add(doc)
        db.commit()
        logger.info("✅ حُفظ في DB: %s", file_id)
    except Exception as e:
        logger.error("فشل الحفظ في DB: %s", e)
        db.rollback()


def _error_result(file_id: str, stored_path, error: str) -> Dict:
    logger.error("❌ فشل %s: %s", file_id, error)
    return {
        "file_id":      file_id,
        "status":       "failed",
        "error":        error,
        "storage_path": str(stored_path) if stored_path else None,
    }
